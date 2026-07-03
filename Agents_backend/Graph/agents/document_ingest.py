"""
document_ingest.py — Parse user-uploaded documents (PDF / DOCX / TXT / MD)
and turn them into `EvidenceItem`s that the existing orchestrator + workers
can cite.

Public surface:
    parse_document(path)            -> {"text", "pages", "metadata"}
    chunk_text(pages, ...)          -> List[Chunk]
    extract_evidence_from_chunks(...) -> List[EvidenceItem]
    derive_topic_from_document(text)-> str
    document_ingest_node(state)     -> dict   (LangGraph node)

The pipeline writes an `evidence.json` file to `uploads/<upload_id>/` at
upload time so the LangGraph node only needs to load it (no re-parsing,
no re-LLM-extraction during the actual job run).
"""

from __future__ import annotations

import json
import logging
import os
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional
import numpy as np

from langchain_core.messages import HumanMessage, SystemMessage

from Graph.state import EvidenceItem, EvidencePack, State
from Graph.templates import RESEARCH_SYSTEM

from .utils import _emit, _job, llm, logger


# ---------------------------------------------------------------------------
# Tunables
# ---------------------------------------------------------------------------

# Hard cap on how many chunks we'll send to the LLM for a single document.
# Keeps cost predictable on a 200-page PDF.
_MAX_CHUNKS = 30

# Default chunking parameters (in characters, not tokens).
# ~6000 chars ≈ 1.5k tokens — comfortable for a single LLM extraction call.
_DEFAULT_CHUNK_CHARS = 6000
_DEFAULT_OVERLAP = 400

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md"}


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class Chunk:
    """A contiguous slice of document text plus the page range it spans."""

    text: str
    page_start: int
    page_end: int


# ---------------------------------------------------------------------------
# 1. PARSING
# ---------------------------------------------------------------------------


def _parse_pdf(path: Path) -> dict:
    """Extract per-page text from a PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "pypdf is required to parse PDF files. Install with `pip install pypdf`."
        ) from e

    reader = PdfReader(str(path))
    pages: List[str] = []
    for page in reader.pages:
        try:
            pages.append((page.extract_text() or "").strip())
        except Exception as exc:  # noqa: BLE001
            logger.warning(f"PDF page extraction failed: {exc}")
            pages.append("")

    text = "\n\n".join(p for p in pages if p)
    return {
        "text": text,
        "pages": pages,
        "metadata": {"format": "pdf", "page_count": len(pages)},
    }


def _parse_docx(path: Path) -> dict:
    """Extract paragraph text from a DOCX. We treat each Heading-1 block
    as a logical 'page' for citation purposes."""
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required to parse DOCX files."
        ) from e

    document = docx.Document(str(path))
    pages: List[str] = []
    buffer: List[str] = []

    for para in document.paragraphs:
        style_name = (para.style.name or "").lower() if para.style else ""
        if style_name.startswith("heading 1") and buffer:
            pages.append("\n".join(buffer).strip())
            buffer = []
        if para.text:
            buffer.append(para.text)

    if buffer:
        pages.append("\n".join(buffer).strip())

    # Fallback: if the doc had no Heading-1 sections, treat the whole thing
    # as a single page.
    if not pages:
        pages = ["\n".join(p.text for p in document.paragraphs if p.text).strip()]

    text = "\n\n".join(p for p in pages if p)
    return {
        "text": text,
        "pages": pages,
        "metadata": {"format": "docx", "page_count": len(pages)},
    }


def _parse_text(path: Path, fmt: str) -> dict:
    """Read a plain text or markdown file."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    # Split on form-feeds (rare) or huge blank gaps, otherwise treat as one page.
    parts = raw.split("\f")
    pages = [p.strip() for p in parts if p.strip()] or [raw.strip()]
    return {
        "text": raw.strip(),
        "pages": pages,
        "metadata": {"format": fmt, "page_count": len(pages)},
    }


def parse_document(path: str | Path) -> dict:
    """Dispatch to the right parser based on file extension.

    Returns a dict shaped:
        {"text": str, "pages": List[str], "metadata": {...}}
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Upload not found: {p}")

    ext = p.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(p)
    if ext == ".docx":
        return _parse_docx(p)
    if ext == ".txt":
        return _parse_text(p, "txt")
    if ext == ".md":
        return _parse_text(p, "md")

    raise ValueError(
        f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTS)}"
    )


# ---------------------------------------------------------------------------
# 2. CHUNKING
# ---------------------------------------------------------------------------


def chunk_text(
    pages: List[str],
    max_chars: int = _DEFAULT_CHUNK_CHARS,
    overlap: int = _DEFAULT_OVERLAP,
) -> List[Chunk]:
    """Greedy chunker that respects page boundaries.

    Pages are concatenated until adding the next page would exceed
    `max_chars`. When a single page is larger than `max_chars`, that page
    is split into multiple chunks of `max_chars - overlap` step.
    """
    chunks: List[Chunk] = []
    if not pages:
        return chunks

    buf = ""
    buf_start_page = 0
    last_page_in_buf = 0

    for page_idx, page in enumerate(pages):
        page = page or ""

        # If a single page is gigantic, hard-split it.
        if len(page) > max_chars:
            # Flush current buffer first.
            if buf:
                chunks.append(Chunk(buf.strip(), buf_start_page + 1, last_page_in_buf + 1))
                buf = ""

            step = max(1, max_chars - overlap)
            for i in range(0, len(page), step):
                segment = page[i : i + max_chars]
                if not segment.strip():
                    continue
                chunks.append(Chunk(segment.strip(), page_idx + 1, page_idx + 1))
            buf_start_page = page_idx
            last_page_in_buf = page_idx
            continue

        # Normal flow: try to extend buffer.
        if not buf:
            buf = page
            buf_start_page = page_idx
            last_page_in_buf = page_idx
            continue

        if len(buf) + 2 + len(page) <= max_chars:
            buf = f"{buf}\n\n{page}"
            last_page_in_buf = page_idx
        else:
            chunks.append(Chunk(buf.strip(), buf_start_page + 1, last_page_in_buf + 1))
            buf = page
            buf_start_page = page_idx
            last_page_in_buf = page_idx

    if buf.strip():
        chunks.append(Chunk(buf.strip(), buf_start_page + 1, last_page_in_buf + 1))

    return chunks


# ---------------------------------------------------------------------------
# 3. EVIDENCE EXTRACTION
# ---------------------------------------------------------------------------


def _extract_one_chunk(chunk: Chunk, filename: str, topic_hint: str) -> List[EvidenceItem]:
    """Send a single chunk to the LLM and return EvidenceItems."""
    extractor = llm.with_structured_output(EvidencePack)

    span = (
        f"page {chunk.page_start}"
        if chunk.page_start == chunk.page_end
        else f"pages {chunk.page_start}-{chunk.page_end}"
    )

    prompt = (
        f"Source document: {filename} ({span})\n"
        f"Working topic: {topic_hint or 'N/A'}\n\n"
        f"Read the following excerpt and extract 3-6 hard, verifiable facts, "
        f"statistics, or specific claims that a writer could cite verbatim. "
        f"Each fact must be self-contained and quotable — no opinions or filler.\n\n"
        f"EXCERPT:\n{chunk.text}"
    )

    try:
        pack = extractor.invoke(
            [
                SystemMessage(content=RESEARCH_SYSTEM),
                HumanMessage(content=prompt),
            ]
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"Doc-evidence extraction failed for {filename} {span}: {exc}")
        return []

    cite_url = f"file://{filename}#p{chunk.page_start}"
    items: List[EvidenceItem] = []
    for ev in pack.evidence:
        items.append(
            EvidenceItem(
                title=ev.title or f"{filename} ({span})",
                url=cite_url,
                snippet=ev.snippet,
                published_at=ev.published_at,
                source=filename,
                authors=ev.authors,
            )
        )
    return items


def extract_evidence_from_chunks(
    chunks: List[Chunk],
    filename: str,
    topic_hint: str = "",
    max_chunks: int = _MAX_CHUNKS,
) -> tuple[List[EvidenceItem], bool]:
    """Run LLM extraction in parallel across chunks.

    Returns (evidence_items, was_truncated). Caps at `max_chunks` to keep
    costs predictable on long documents.
    """
    if not chunks:
        return [], False

    truncated = len(chunks) > max_chunks
    work = chunks[:max_chunks]

    out: List[EvidenceItem] = []
    if not work:
        return out, truncated

    with ThreadPoolExecutor(max_workers=4) as ex:
        futures = {
            ex.submit(_extract_one_chunk, c, filename, topic_hint): c for c in work
        }
        for fut in as_completed(futures):
            try:
                out.extend(fut.result() or [])
            except Exception as exc:  # noqa: BLE001
                logger.warning(f"Chunk extractor raised: {exc}")

    # De-duplicate by snippet (LLM occasionally repeats stats across chunks).
    seen: set[str] = set()
    deduped: List[EvidenceItem] = []
    for item in out:
        key = (item.snippet or "").strip().lower()[:160]
        if key and key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped, truncated


# ---------------------------------------------------------------------------
# 4. TOPIC DERIVATION (used only when source_mode == "auto_topic")
# ---------------------------------------------------------------------------


def derive_topic_from_document(text: str) -> str:
    """Return a 6-10 word working title derived from the document body."""
    excerpt = (text or "").strip()
    if not excerpt:
        return ""
    excerpt = excerpt[:6000]

    try:
        resp = llm.invoke(
            [
                SystemMessage(
                    content=(
                        "You write short, SEO-friendly working titles. "
                        "Read the document excerpt and return ONE line of 6-10 words "
                        "that captures its core subject. No quotes, no punctuation at the end."
                    )
                ),
                HumanMessage(content=f"DOCUMENT EXCERPT:\n{excerpt}"),
            ]
        )
        title = (resp.content or "").strip().splitlines()[0].strip(" \"'.")
        return title[:120]
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"derive_topic_from_document failed: {exc}")
        return ""


# ---------------------------------------------------------------------------
# 5. UPLOAD HELPER (used by api.py at upload time)
# ---------------------------------------------------------------------------


def split_into_sentences(text: str) -> List[str]:
    """A simple but effective sentence splitter using regex."""
    sentence_endings = re.compile(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|\!)\s')
    sentences = sentence_endings.split(text)
    return [s.strip() for s in sentences if s.strip()]


def get_embeddings_model():
    """Dynamically initializes the OpenAI Embeddings model to avoid startup errors if key is missing."""
    from langchain_openai import OpenAIEmbeddings
    return OpenAIEmbeddings(model="text-embedding-3-small")


def semantic_chunking(
    pages: List[str],
    embeddings_model,
    target_chunk_size: int = 1500,
    min_chunk_size: int = 300,
    max_chunk_size: int = 6000,
    distance_threshold_percentile: float = 85.0
) -> List[Chunk]:
    """
    Groups sentences into semantically coherent chunks using sentence embeddings.
    """
    # 1. Break the document down into sentences, tracking which page they came from
    sentence_records = []
    for page_idx, page in enumerate(pages):
        if not page:
            continue
        sentences = split_into_sentences(page)
        for s in sentences:
            sentence_records.append({
                "text": s,
                "page": page_idx + 1
            })
            
    if not sentence_records:
        return []

    # 2. Batch embed the sentences
    texts = [record["text"] for record in sentence_records]
    
    # Batch embedding API (OpenAI embeddings model)
    try:
        embeddings = embeddings_model.embed_documents(texts)
    except Exception as exc:
        logger.warning(f"Embedding generation failed: {exc}. Falling back to character-based chunking.")
        # Fallback to simple character-based chunking if embedding fails
        return chunk_text(pages, max_chars=max_chunk_size)

    # 3. Calculate distance between consecutive sentences
    distances = []
    for i in range(len(embeddings) - 1):
        vec1 = np.array(embeddings[i])
        vec2 = np.array(embeddings[i+1])
        # Cosine distance
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        if norm1 > 0 and norm2 > 0:
            cos_sim = np.dot(vec1, vec2) / (norm1 * norm2)
        else:
            cos_sim = 1.0
        distances.append(1.0 - cos_sim)

    # 4. Determine splitting threshold
    if distances:
        threshold = np.percentile(distances, distance_threshold_percentile)
    else:
        threshold = 1.0

    # 5. Build chunks based on threshold, respecting target/min/max sizes
    chunks: List[Chunk] = []
    current_sentences = []
    current_length = 0
    start_page = sentence_records[0]["page"]
    
    for idx, record in enumerate(sentence_records):
        current_sentences.append(record["text"])
        current_length += len(record["text"]) + 1 # +1 for space/newline
        
        # Determine if we should split before the NEXT sentence
        should_split = False
        if idx < len(sentence_records) - 1:
            # Split if semantic distance exceeds threshold
            is_semantic_boundary = distances[idx] > threshold
            # Avoid making chunks too small or too large
            if is_semantic_boundary and current_length >= min_chunk_size:
                should_split = True
            if current_length >= target_chunk_size:
                should_split = True
        
        # Hard cap ceiling limit
        if current_length >= max_chunk_size:
            should_split = True
            
        if should_split or idx == len(sentence_records) - 1:
            chunk_text_str = " ".join(current_sentences).strip()
            if chunk_text_str:
                chunks.append(Chunk(
                    text=chunk_text_str,
                    page_start=start_page,
                    page_end=record["page"]
                ))
            current_sentences = []
            current_length = 0
            if idx < len(sentence_records) - 1:
                start_page = sentence_records[idx + 1]["page"]

    return chunks


def ingest_upload(upload_dir: Path, original_filename: str) -> dict:
    """End-to-end upload processing: parse → chunk → extract → persist.
    
    Now supports Advanced RAG: semantic chunking and chunk embeddings.
    """
    src_files = [
        p for p in upload_dir.iterdir()
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS
    ]
    if not src_files:
        raise FileNotFoundError(f"No supported source file in {upload_dir}")
    src = src_files[0]

    parsed = parse_document(src)
    
    # ── Semantic Chunking ───────────────────────────────────────────
    try:
        embeddings_model = get_embeddings_model()
        chunks = semantic_chunking(parsed["pages"], embeddings_model)
    except Exception as exc:
        logger.warning(f"Semantic chunking failed: {exc}. Using character-based fallback.")
        chunks = chunk_text(parsed["pages"])

    # ── Chunk Embeddings ─────────────────────────────────────────────
    chunk_embeddings = []
    if chunks:
        try:
            chunk_texts = [c.text for c in chunks]
            chunk_embeddings = embeddings_model.embed_documents(chunk_texts)
        except Exception as exc:
            logger.warning(f"Generating chunk embeddings failed: {exc}")

    # Save embeddings & chunks
    embeddings_path = upload_dir / "embeddings.json"
    with embeddings_path.open("w", encoding="utf-8") as f:
        serialized_chunks = []
        for i, c in enumerate(chunks):
            serialized_chunks.append({
                "text": c.text,
                "page_start": c.page_start,
                "page_end": c.page_end,
                "embedding": chunk_embeddings[i] if i < len(chunk_embeddings) else []
            })
        json.dump(serialized_chunks, f, indent=2)

    # ── Evidence Extraction (Fallback / General) ────────────────────
    evidence, truncated = extract_evidence_from_chunks(
        chunks, original_filename, topic_hint=""
    )
    derived_topic = derive_topic_from_document(parsed["text"])

    # Persist evidence so document_ingest_node can load it cheaply as fallback
    evidence_path = upload_dir / "evidence.json"
    with evidence_path.open("w", encoding="utf-8") as f:
        json.dump([e.model_dump() for e in evidence], f, indent=2)

    preview = (parsed["text"] or "")[:600]

    meta = {
        "filename": original_filename,
        "stored_path": str(src),
        "format": parsed["metadata"].get("format"),
        "pages": parsed["metadata"].get("page_count", 0),
        "chunks": len(chunks),
        "chunks_processed": min(len(chunks), _MAX_CHUNKS),
        "evidence_count": len(evidence),
        "truncated": truncated,
        "derived_topic": derived_topic,
        "preview": preview,
    }
    with (upload_dir / "meta.json").open("w", encoding="utf-8") as f:
        json.dump(meta, f, indent=2)

    return meta


def load_upload_metadata(upload_dir: Path) -> Optional[dict]:
    meta_file = upload_dir / "meta.json"
    if not meta_file.exists():
        return None
    try:
        return json.loads(meta_file.read_text(encoding="utf-8"))
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"Failed to read upload meta {meta_file}: {exc}")
        return None


def load_upload_evidence(upload_dir: Path) -> List[EvidenceItem]:
    ev_file = upload_dir / "evidence.json"
    if not ev_file.exists():
        return []
    try:
        raw = json.loads(ev_file.read_text(encoding="utf-8"))
        return [EvidenceItem(**item) for item in raw]
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"Failed to load doc evidence {ev_file}: {exc}")
        return []


# ---------------------------------------------------------------------------
# 6. PIPELINE NODE
# ---------------------------------------------------------------------------


def _resolve_uploads_root() -> Path:
    """Same root that api.py uses (Agents_backend/uploads/)."""
    here = Path(__file__).resolve()
    # …/Agents_backend/Graph/agents/document_ingest.py → up 3 = Agents_backend
    backend_dir = here.parents[2]
    root = backend_dir / "uploads"
    root.mkdir(parents=True, exist_ok=True)
    return root


def document_ingest_node(state: State) -> dict:
    """LangGraph node: load doc evidence + (optionally) derive topic.

    Runs only when `state["upload_id"]` is set.
    Performs dynamic Advanced RAG search over embedded chunks if topic or keywords are specified,
    falling back to pre-extracted evidence if embeddings do not exist or retrieval fails.
    """
    job_id = _job(state)
    upload_id = state.get("upload_id") or ""
    source_mode = state.get("source_mode") or "hybrid"
    topic = state.get("topic") or ""
    keywords = state.get("target_keywords") or []

    if not upload_id:
        return {}

    _emit(job_id, "ingest", "started",
          "Performing Advanced RAG semantic search on document...",
          {"upload_id": upload_id, "source_mode": source_mode})
    logger.info(f"📄 INGEST --- upload_id={upload_id} mode={source_mode}")

    upload_dir = _resolve_uploads_root() / upload_id
    if not upload_dir.exists():
        _emit(job_id, "ingest", "error", f"Upload {upload_id} not found.")
        return {}

    meta = load_upload_metadata(upload_dir) or {}
    filename = meta.get("filename", "document")
    
    # ── Advanced RAG Retrieval ──────────────────────────────────────
    evidence = []
    retrieved_chunks = []
    
    embeddings_path = upload_dir / "embeddings.json"
    if embeddings_path.exists() and (topic or keywords):
        try:
            with embeddings_path.open("r", encoding="utf-8") as f:
                serialized_chunks = json.load(f)
            
            # Reconstruct Chunk objects
            all_chunks = []
            chunk_embs = []
            for item in serialized_chunks:
                all_chunks.append(Chunk(
                    text=item["text"],
                    page_start=item["page_start"],
                    page_end=item["page_end"]
                ))
                chunk_embs.append(item.get("embedding", []))
                
            # Perform query embedding
            query_text = f"{topic} " + " ".join(keywords)
            query_text = query_text.strip()
            
            if query_text and chunk_embs and any(chunk_embs):
                embeddings_model = get_embeddings_model()
                query_emb = embeddings_model.embed_query(query_text)
                
                # Compute cosine similarities
                similarities = []
                for emb in chunk_embs:
                    if not emb:
                        similarities.append(0.0)
                        continue
                    v_q = np.array(query_emb)
                    v_c = np.array(emb)
                    norm_q = np.linalg.norm(v_q)
                    norm_c = np.linalg.norm(v_c)
                    if norm_q > 0 and norm_c > 0:
                        sim = np.dot(v_q, v_c) / (norm_q * norm_c)
                    else:
                        sim = 0.0
                    similarities.append(sim)
                    
                # Rank chunks and filter by similarity threshold (0.35)
                ranked_indices = np.argsort(similarities)[::-1]
                top_indices = [idx for idx in ranked_indices if similarities[idx] >= 0.35]
                
                # If nothing passed the threshold, fall back to top 2 chunks (if similarity > 0.0)
                if not top_indices:
                    top_indices = [idx for idx in ranked_indices if similarities[idx] > 0.0][:2]
                else:
                    top_indices = top_indices[:8]
                    
                retrieved_chunks = [all_chunks[idx] for idx in top_indices]
                
                logger.info(f"RAG: Retrieved top {len(retrieved_chunks)} semantic chunks for query '{query_text}'")
                _emit(job_id, "ingest", "working", 
                      f"Retrieved top {len(retrieved_chunks)} relevant sections from document...")
        except Exception as exc:
            logger.exception(f"Advanced RAG retrieval failed, falling back to pre-extracted evidence: {exc}")
            
    # If we retrieved chunks, perform dynamic evidence extraction
    if retrieved_chunks:
        try:
            evidence, _ = extract_evidence_from_chunks(
                retrieved_chunks, filename, topic_hint=topic, max_chunks=8
            )
        except Exception as exc:
            logger.exception(f"Dynamic evidence extraction failed: {exc}")

    # Fallback to pre-extracted evidence.json if dynamic RAG yielded nothing
    if not evidence:
        logger.info("Using pre-extracted evidence fallback")
        evidence = load_upload_evidence(upload_dir)

    out: dict = {
        "evidence": evidence,
        "document_filename": filename,
    }

    # Topic auto-derivation
    if source_mode == "auto_topic":
        derived = (meta.get("derived_topic") or "").strip()
        topic = (state.get("topic") or "").strip()
        if derived and not topic:
            out["topic"] = derived

    _emit(
        job_id, "ingest", "completed",
        f"Retrieved and extracted {len(evidence)} evidence item(s) from {filename}",
        {
            "evidence_count": len(evidence),
            "pages": meta.get("pages", 0),
            "chunks": meta.get("chunks", len(retrieved_chunks) or meta.get("chunks", 0)),
            "truncated": meta.get("truncated", False),
        },
    )
    logger.info(f"✅ Ingest complete: {len(evidence)} evidence item(s)")
    return out
